"""
═══════════════════════════════════════════════════════════════════════════════
  Compare Pipeline (All Methods: Original, P1, P3, P4 Ensemble)
  ─────────────────────────────────────────────────────────────────
  Runs all model versions on a video to compare how augmented training
  improves predictions.

  Usage:
      python priority_3_augmentation/compare_pipeline_augmented.py --video path/to/video.mp4
═══════════════════════════════════════════════════════════════════════════════
"""

import cv2
import torch
import torch.nn as nn
import numpy as np
from torchvision import transforms, models
from ultralytics import YOLO
from pathlib import Path
from PIL import Image
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random
import argparse
import datetime
import sys

sys.path.insert(0, str(Path(__file__).parent.parent))
from sperm_pipeline.pipeline import SpermAnalysisPipeline

# ═════════════════════════════════════════════════════════════════════════════
# CONFIGURATION
# ═════════════════════════════════════════════════════════════════════════════
BASE_DIR = Path(r"D:\Krishna\SMDMSSdataset")
DETECTION_MODEL = BASE_DIR / "sperm.v1i.yolov8" / "runs" / "detect" / "sperm_detection" / "weights" / "best.pt"

# All model variants
MODELS = {
    "YOLO_Orig": BASE_DIR / "runs" / "classify" / "runs" / "sperm_cls_balanced" / "weights" / "best.pt",
    "YOLO_P1": BASE_DIR / "priority_1_class_weights" / "yolo_runs" / "sperm_cls_weighted" / "weights" / "best.pt",
    "YOLO_P3": BASE_DIR / "priority_3_augmentation" / "yolo_runs" / "sperm_cls_augmented" / "weights" / "best.pt",
    "EffNet_Orig": BASE_DIR / "efficientnet_training" / "weights" / "efficientnet_b0_best.pt",
    "EffNet_P1": BASE_DIR / "priority_1_class_weights" / "weights" / "efficientnet_b0_weighted.pt",
    "EffNet_P3": BASE_DIR / "priority_3_augmentation" / "weights" / "efficientnet_b0_augmented.pt",
}

IMG_SIZE = 224
CLASS_NAMES = ["Combined_Anomaly", "Head_Anomaly", "Midpiece_Anomaly", "Normal", "Tail_Anomaly"]
OUTPUT_DOCX = BASE_DIR / "priority_3_augmentation" / "augmented_comparison_report.docx"

# P4 Ensemble weights
YOLO_WEIGHT = 0.4
EFFNET_WEIGHT = 0.6


def load_effnet(path, device):
    model = models.efficientnet_b0(weights=None)
    num_ftrs = model.classifier[1].in_features
    try:
        model.classifier = nn.Sequential(nn.Dropout(p=0.4, inplace=True), nn.Linear(num_ftrs, 5))
        model.load_state_dict(torch.load(path, map_location=device))
    except RuntimeError:
        model = models.efficientnet_b0(weights=None)
        model.classifier = nn.Sequential(nn.Dropout(p=0.35, inplace=True), nn.Linear(model.classifier[1].in_features, 5))
        try:
            model.load_state_dict(torch.load(path, map_location=device))
        except RuntimeError:
            model = models.efficientnet_b0(weights=None)
            model.classifier[1] = nn.Linear(model.classifier[1].in_features, 5)
            model.load_state_dict(torch.load(path, map_location=device))
    model.to(device).eval()
    return model


def classify_effnet(model, crop, device):
    tfm = transforms.Compose([
        transforms.Resize((IMG_SIZE, IMG_SIZE)), transforms.ToTensor(),
        transforms.Normalize([0.485, 0.456, 0.406], [0.229, 0.224, 0.225])
    ])
    img = Image.fromarray(cv2.cvtColor(crop, cv2.COLOR_BGR2RGB))
    with torch.no_grad():
        out = model(tfm(img).unsqueeze(0).to(device))
        probs = torch.nn.functional.softmax(out, dim=1)[0]
    idx = probs.argmax().item()
    return CLASS_NAMES[idx], float(probs[idx])


def get_effnet_probs(model, crop, device):
    """Get full probability vector from EfficientNet."""
    tfm = transforms.Compose([
        transforms.Resize((IMG_SIZE, IMG_SIZE)), transforms.ToTensor(),
        transforms.Normalize([0.485, 0.456, 0.406], [0.229, 0.224, 0.225])
    ])
    img = Image.fromarray(cv2.cvtColor(crop, cv2.COLOR_BGR2RGB))
    with torch.no_grad():
        out = model(tfm(img).unsqueeze(0).to(device))
        probs = torch.nn.functional.softmax(out, dim=1)[0].cpu().numpy()
    return probs


def get_yolo_probs(model, crop):
    """Get full probability vector from YOLOv8 (remapped to CLASS_NAMES order)."""
    res = model.predict(cv2.resize(crop, (IMG_SIZE, IMG_SIZE)), imgsz=IMG_SIZE, verbose=False)
    probs_tensor = res[0].probs.data.cpu().numpy()
    yolo_names = res[0].names
    remapped = np.zeros(len(CLASS_NAMES))
    for yolo_idx, name in yolo_names.items():
        if name in CLASS_NAMES:
            remapped[CLASS_NAMES.index(name)] = probs_tensor[yolo_idx]
    return remapped


def classify_yolo(model, crop):
    res = model.predict(cv2.resize(crop, (IMG_SIZE, IMG_SIZE)), imgsz=IMG_SIZE, verbose=False)
    return res[0].names[res[0].probs.top1], float(res[0].probs.top1conf)


def extract_frames(video_path, n=2):
    cap = cv2.VideoCapture(str(video_path))
    total = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    indices = sorted(random.sample(range(total), min(n, total)))
    frames = []
    for idx in indices:
        cap.set(cv2.CAP_PROP_POS_FRAMES, idx)
        ret, frame = cap.read()
        if ret: frames.append((idx, frame))
    cap.release()
    return frames


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--video", required=True)
    parser.add_argument("--frames", type=int, default=2)
    args = parser.parse_args()

    device = torch.device("cuda:0" if torch.cuda.is_available() else "cpu")
    print("═" * 65)
    print("  All Methods Comparison: Original, P1, P3, P4 Ensemble")
    print("═" * 65)
    print(f"  Device: {device}\n")

    # Load detection pipeline
    det_cls = MODELS["YOLO_Orig"] if MODELS["YOLO_Orig"].exists() else MODELS["YOLO_P1"]
    pipeline = SpermAnalysisPipeline(str(DETECTION_MODEL), str(det_cls))
    pipeline.load_models()

    # Load available models
    loaded = {}
    for name, path in MODELS.items():
        if not path.exists():
            print(f"  ⚠️ {name}: not found, skipping")
            continue
        if name.startswith("YOLO"):
            loaded[name] = ("yolo", YOLO(str(path)))
        else:
            loaded[name] = ("effnet", load_effnet(path, device))
        print(f"  ✅ {name}: loaded")

    frames = extract_frames(args.video, args.frames)
    print(f"\n🎥 Extracted {len(frames)} frames\n")

    # Check if P4 ensemble is possible (needs YOLO_P1 + EffNet_P1)
    can_ensemble = "YOLO_P1" in loaded and "EffNet_P1" in loaded
    if can_ensemble:
        print("  ✅ Ensemble_P4: enabled (YOLO_P1 + EffNet_P1)")
    else:
        print("  ⚠️ Ensemble_P4: disabled (needs both YOLO_P1 and EffNet_P1)")

    results = []
    all_model_names = list(loaded.keys()) + (["Ensemble_P4"] if can_ensemble else [])

    for fidx, frame in frames:
        dets = pipeline.detect_sperm(frame)
        print(f"── Frame {fidx}: {len(dets)} sperm ──")
        for sid, det in enumerate(dets, 1):
            crop = pipeline.crop_detection(frame, det["bbox"])
            if crop.size == 0: continue
            row = {"frame": fidx, "sperm": sid}
            preds = []
            for mname, (mtype, model) in loaded.items():
                if mtype == "yolo":
                    p, c = classify_yolo(model, crop)
                else:
                    p, c = classify_effnet(model, crop, device)
                row[f"{mname}_pred"] = p
                row[f"{mname}_conf"] = c
                preds.append(f"{mname}={p:20s}")

            # P4 Ensemble: average P1 YOLO + P1 EffNet probability vectors
            if can_ensemble:
                yolo_probs = get_yolo_probs(loaded["YOLO_P1"][1], crop)
                effnet_probs = get_effnet_probs(loaded["EffNet_P1"][1], crop, device)
                combined = YOLO_WEIGHT * yolo_probs + EFFNET_WEIGHT * effnet_probs
                combined = combined / combined.sum()
                ens_idx = np.argmax(combined)
                row["Ensemble_P4_pred"] = CLASS_NAMES[ens_idx]
                row["Ensemble_P4_conf"] = float(combined[ens_idx])
                preds.append(f"Ensemble_P4={CLASS_NAMES[ens_idx]:20s}")

            print(f"  #{sid:2d} " + " | ".join(preds))
            results.append(row)

    # Generate report
    doc = Document()
    doc.add_heading("All Methods Comparison Report", level=0)
    doc.add_paragraph(f"Generated: {datetime.datetime.now():%Y-%m-%d %H:%M}")
    doc.add_paragraph(f"Video: {Path(args.video).name} | Crops: {len(results)}")
    if can_ensemble:
        doc.add_paragraph(f"P4 Ensemble: {YOLO_WEIGHT:.0%} YOLO_P1 + {EFFNET_WEIGHT:.0%} EffNet_P1")

    # Class distribution table
    doc.add_heading("Class Distribution by Model", level=1)
    tbl = doc.add_table(rows=len(CLASS_NAMES)+1, cols=len(all_model_names)+1)
    tbl.style = "Light Grid Accent 1"
    tbl.rows[0].cells[0].text = "Class"
    for j, mn in enumerate(all_model_names, 1):
        tbl.rows[0].cells[j].text = mn
    for i, cls in enumerate(CLASS_NAMES, 1):
        tbl.rows[i].cells[0].text = cls
        for j, mn in enumerate(all_model_names, 1):
            cnt = sum(1 for r in results if r.get(f"{mn}_pred") == cls)
            tbl.rows[i].cells[j].text = f"{cnt} ({cnt/max(len(results),1):.0%})"

    doc.save(str(OUTPUT_DOCX))
    print(f"\n✅ Report: {OUTPUT_DOCX}")

    # Summary
    total = len(results)
    print(f"\n{'═'*65}")
    for mn in all_model_names:
        print(f"  {mn}:")
        for cls in CLASS_NAMES:
            cnt = sum(1 for r in results if r.get(f"{mn}_pred") == cls)
            print(f"    {cls:25s} {cnt:3d} ({cnt/max(total,1):.0%})")
    print(f"{'═'*65}")


if __name__ == "__main__":
    main()

