"""
═══════════════════════════════════════════════════════════════════════════════
  Full-Video Analysis: All Models × All Frames
  ─────────────────────────────────────────────
  Processes EVERY Nth frame of the video through ALL 7 models:
    - YOLO Original, YOLO P1, YOLO P3
    - EffNet Original, EffNet P1, EffNet P3
    - Ensemble P4 (40% YOLO_P1 + 60% EffNet_P1)

  Computes per-frame and overall average class distributions.

  Usage:
      python full_video_analysis.py --video path/to/video.mp4
      python full_video_analysis.py --video path/to/video.mp4 --every 25
═══════════════════════════════════════════════════════════════════════════════
"""

import cv2
import torch
import torch.nn as nn
import numpy as np
from torchvision import transforms, models
from ultralytics import YOLO
from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import argparse
import datetime
import sys
import time

sys.path.insert(0, str(Path(__file__).parent))
from sperm_pipeline.pipeline import SpermAnalysisPipeline

# ═════════════════════════════════════════════════════════════════════════════
# CONFIGURATION
# ═════════════════════════════════════════════════════════════════════════════
BASE_DIR = Path(r"D:\Krishna\SMDMSSdataset")
DETECTION_MODEL = BASE_DIR / "sperm.v1i.yolov8" / "runs" / "detect" / "sperm_detection" / "weights" / "best.pt"

MODELS = {
    "YOLO_Orig": BASE_DIR / "runs" / "classify" / "runs" / "sperm_cls_balanced" / "weights" / "best.pt",
    "YOLO_P1":   BASE_DIR / "priority_1_class_weights" / "yolo_runs" / "sperm_cls_weighted" / "weights" / "best.pt",
    "YOLO_P3":   BASE_DIR / "priority_3_augmentation" / "yolo_runs" / "sperm_cls_augmented" / "weights" / "best.pt",
    "EffNet_Orig": BASE_DIR / "efficientnet_training" / "weights" / "efficientnet_b0_best.pt",
    "EffNet_P1":   BASE_DIR / "priority_1_class_weights" / "weights" / "efficientnet_b0_weighted.pt",
    "EffNet_P3":   BASE_DIR / "priority_3_augmentation" / "weights" / "efficientnet_b0_augmented.pt",
}

IMG_SIZE = 224
CLASS_NAMES = ["Combined_Anomaly", "Head_Anomaly", "Midpiece_Anomaly", "Normal", "Tail_Anomaly"]

# P4 Ensemble
YOLO_W = 0.4
EFFNET_W = 0.6


# ═════════════════════════════════════════════════════════════════════════════
# MODEL LOADING
# ═════════════════════════════════════════════════════════════════════════════

def load_effnet(path, device):
    model = models.efficientnet_b0(weights=None)
    num_ftrs = model.classifier[1].in_features
    for dropout_val in [0.4, 0.35]:
        try:
            model = models.efficientnet_b0(weights=None)
            num_ftrs = model.classifier[1].in_features
            model.classifier = nn.Sequential(
                nn.Dropout(p=dropout_val, inplace=True),
                nn.Linear(num_ftrs, len(CLASS_NAMES))
            )
            model.load_state_dict(torch.load(path, map_location=device))
            model.to(device).eval()
            return model
        except RuntimeError:
            continue
    # Final fallback: no dropout
    model = models.efficientnet_b0(weights=None)
    model.classifier[1] = nn.Linear(model.classifier[1].in_features, len(CLASS_NAMES))
    model.load_state_dict(torch.load(path, map_location=device))
    model.to(device).eval()
    return model


# ═════════════════════════════════════════════════════════════════════════════
# CLASSIFICATION FUNCTIONS
# ═════════════════════════════════════════════════════════════════════════════

_effnet_tfm = transforms.Compose([
    transforms.Resize((IMG_SIZE, IMG_SIZE)),
    transforms.ToTensor(),
    transforms.Normalize([0.485, 0.456, 0.406], [0.229, 0.224, 0.225])
])


def classify_effnet(model, crop, device):
    img = Image.fromarray(cv2.cvtColor(crop, cv2.COLOR_BGR2RGB))
    with torch.no_grad():
        out = model(_effnet_tfm(img).unsqueeze(0).to(device))
        probs = torch.nn.functional.softmax(out, dim=1)[0].cpu().numpy()
    idx = int(np.argmax(probs))
    return CLASS_NAMES[idx], float(probs[idx]), probs


def classify_yolo(model, crop):
    res = model.predict(cv2.resize(crop, (IMG_SIZE, IMG_SIZE)), imgsz=IMG_SIZE, verbose=False)
    probs_raw = res[0].probs.data.cpu().numpy()
    yolo_names = res[0].names
    probs = np.zeros(len(CLASS_NAMES))
    for yi, name in yolo_names.items():
        if name in CLASS_NAMES:
            probs[CLASS_NAMES.index(name)] = probs_raw[yi]
    idx = int(np.argmax(probs))
    return CLASS_NAMES[idx], float(probs[idx]), probs


# ═════════════════════════════════════════════════════════════════════════════
# MAIN
# ═════════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description="Full-Video All-Model Analysis")
    parser.add_argument("--video", required=True, help="Path to input video")
    parser.add_argument("--every", type=int, default=1,
                        help="Process every Nth frame (default: 1 = ALL frames)")
    args = parser.parse_args()

    video_path = Path(args.video)
    if not video_path.exists():
        print(f"❌ Video not found: {video_path}")
        sys.exit(1)

    device = torch.device("cuda:0" if torch.cuda.is_available() else "cpu")

    print("═" * 70)
    print("  FULL VIDEO ANALYSIS — All Models × All Frames")
    print("═" * 70)
    print(f"  Video:  {video_path.name}")
    print(f"  Device: {device}")
    print(f"  Sample: every {args.every}th frame\n")

    # ── Open video ──
    cap = cv2.VideoCapture(str(video_path))
    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    fps = cap.get(cv2.CAP_PROP_FPS)
    frames_to_process = list(range(0, total_frames, args.every))
    print(f"  Total frames: {total_frames} | FPS: {fps:.1f}")
    print(f"  Frames to process: {len(frames_to_process)}\n")

    # ── Load detection pipeline ──
    print("🔬 Loading detection pipeline...")
    det_cls = MODELS["YOLO_Orig"] if MODELS["YOLO_Orig"].exists() else MODELS["YOLO_P1"]
    pipeline = SpermAnalysisPipeline(str(DETECTION_MODEL), str(det_cls))
    pipeline.load_models()

    # ── Load classifiers ──
    print("\n🧠 Loading all classifiers...")
    loaded = {}
    for name, path in MODELS.items():
        if not path.exists():
            print(f"  ⚠️  {name}: NOT FOUND, skipping")
            continue
        if name.startswith("YOLO"):
            loaded[name] = ("yolo", YOLO(str(path)))
        else:
            loaded[name] = ("effnet", load_effnet(path, device))
        print(f"  ✅ {name}")

    can_ensemble = "YOLO_P1" in loaded and "EffNet_P1" in loaded
    all_model_names = list(loaded.keys()) + (["Ensemble_P4"] if can_ensemble else [])
    print(f"\n  Models: {len(all_model_names)} ({', '.join(all_model_names)})")

    # ── Initialize counters ──
    # Per-model totals: model_name → {class_name: count}
    global_counts = {mn: {c: 0 for c in CLASS_NAMES} for mn in all_model_names}
    # Per-frame stats: list of dicts
    frame_stats = []
    total_sperm = 0

    # ── Process frames ──
    start_time = time.time()
    print(f"\n{'─'*70}")
    print(f"  Processing {len(frames_to_process)} frames...")
    print(f"{'─'*70}\n")

    for proc_idx, frame_idx in enumerate(frames_to_process):
        cap.set(cv2.CAP_PROP_POS_FRAMES, frame_idx)
        ret, frame = cap.read()
        if not ret:
            continue

        detections = pipeline.detect_sperm(frame)
        n_sperm = len(detections)

        if n_sperm == 0:
            continue

        # Per-frame class counts
        frame_counts = {mn: {c: 0 for c in CLASS_NAMES} for mn in all_model_names}

        for det in detections:
            crop = pipeline.crop_detection(frame, det["bbox"])
            if crop.size == 0:
                continue

            total_sperm += 1

            # Classify with each model
            yolo_p1_probs = None
            effnet_p1_probs = None

            for mname, (mtype, model) in loaded.items():
                if mtype == "yolo":
                    pred, conf, probs = classify_yolo(model, crop)
                    if mname == "YOLO_P1":
                        yolo_p1_probs = probs
                else:
                    pred, conf, probs = classify_effnet(model, crop, device)
                    if mname == "EffNet_P1":
                        effnet_p1_probs = probs

                global_counts[mname][pred] += 1
                frame_counts[mname][pred] += 1

            # Ensemble P4
            if can_ensemble and yolo_p1_probs is not None and effnet_p1_probs is not None:
                combined = YOLO_W * yolo_p1_probs + EFFNET_W * effnet_p1_probs
                combined = combined / combined.sum()
                ens_pred = CLASS_NAMES[int(np.argmax(combined))]
                global_counts["Ensemble_P4"][ens_pred] += 1
                frame_counts["Ensemble_P4"][ens_pred] += 1

        frame_stats.append({
            "frame_idx": frame_idx,
            "n_sperm": n_sperm,
            "counts": frame_counts,
        })

        # Progress
        if (proc_idx + 1) % 10 == 0 or proc_idx == 0:
            elapsed = time.time() - start_time
            pct = (proc_idx + 1) / len(frames_to_process) * 100
            eta = elapsed / (proc_idx + 1) * (len(frames_to_process) - proc_idx - 1)
            print(f"  [{proc_idx+1:4d}/{len(frames_to_process)}] "
                  f"Frame {frame_idx:6d} | {n_sperm:2d} sperm | "
                  f"{pct:.0f}% done | ETA: {eta:.0f}s")

    cap.release()
    elapsed_total = time.time() - start_time

    # ═════════════════════════════════════════════════════════════════════════
    # RESULTS
    # ═════════════════════════════════════════════════════════════════════════

    print(f"\n{'═'*70}")
    print(f"  RESULTS: {total_sperm} total sperm across {len(frame_stats)} frames")
    print(f"  Time: {elapsed_total:.1f}s ({elapsed_total/max(len(frame_stats),1):.2f}s/frame)")
    print(f"{'═'*70}\n")

    # ── Class distribution per model ──
    print(f"  {'Model':<16s}", end="")
    for c in CLASS_NAMES:
        short = c.replace("_Anomaly", "").replace("Combined", "Comb")[:8]
        print(f" {short:>8s}", end="")
    print(f" {'Anomaly%':>9s}")
    print(f"  {'─'*16}" + "─" * 8 * len(CLASS_NAMES) + "─" * 10)

    for mn in all_model_names:
        counts = global_counts[mn]
        total = sum(counts.values())
        if total == 0:
            continue
        anomaly_pct = (total - counts["Normal"]) / total * 100

        print(f"  {mn:<16s}", end="")
        for c in CLASS_NAMES:
            pct = counts[c] / total * 100
            print(f" {pct:7.1f}%", end="")
        print(f" {anomaly_pct:8.1f}%")

    # Clinical reference
    print(f"\n  {'CLINICAL':<16s}", end="")
    clinical = {"Combined_Anomaly": 0, "Head_Anomaly": 98.0,
                "Midpiece_Anomaly": 23.3, "Normal": 2.0, "Tail_Anomaly": 15.5}
    for c in CLASS_NAMES:
        print(f" {clinical[c]:7.1f}%", end="")
    print(f" {98.0:8.1f}%")

    # ═════════════════════════════════════════════════════════════════════════
    # MAE (Mean Absolute Error) — per model vs clinical
    # ═════════════════════════════════════════════════════════════════════════
    print(f"\n{'═'*70}")
    print(f"  MAE (Mean Absolute Error) vs Clinical Ground Truth")
    print(f"{'═'*70}\n")

    # Header
    print(f"  {'Model':<16s}", end="")
    for c in CLASS_NAMES:
        short = c.replace("_Anomaly", "").replace("Combined", "Comb")[:8]
        print(f" {short:>8s}", end="")
    print(f" {'MAE':>8s}")
    print(f"  {'─'*16}" + "─" * 8 * len(CLASS_NAMES) + "─" * 9)

    mae_results = {}
    for mn in all_model_names:
        counts = global_counts[mn]
        total = sum(counts.values())
        if total == 0:
            continue

        errors = []
        print(f"  {mn:<16s}", end="")
        for c in CLASS_NAMES:
            predicted_pct = counts[c] / total * 100
            clinical_pct = clinical[c]
            abs_err = abs(predicted_pct - clinical_pct)
            errors.append(abs_err)
            print(f" {abs_err:7.1f}%", end="")

        mae = np.mean(errors)
        mae_results[mn] = mae
        print(f" {mae:7.1f}%")

    # Rank models by MAE
    print(f"\n  🏆 Model Ranking (lowest MAE = best):")
    ranked = sorted(mae_results.items(), key=lambda x: x[1])
    for rank, (mn, mae) in enumerate(ranked, 1):
        medal = {1: "🥇", 2: "🥈", 3: "🥉"}.get(rank, "  ")
        print(f"   {medal} {rank}. {mn:<16s}  MAE = {mae:.2f}%")

    print(f"{'═'*70}")

    # ── Generate DOCX report ──
    output_docx = BASE_DIR / "full_video_analysis_report.docx"
    doc = Document()
    doc.add_heading("Full Video Analysis — All Models", level=0)
    doc.add_paragraph(f"Generated: {datetime.datetime.now():%Y-%m-%d %H:%M:%S}")
    doc.add_paragraph(f"Video: {video_path.name}")
    doc.add_paragraph(f"Total frames processed: {len(frame_stats)} (every {args.every}th of {total_frames})")
    doc.add_paragraph(f"Total sperm classified: {total_sperm}")
    doc.add_paragraph(f"Processing time: {elapsed_total:.1f}s")

    # Summary table
    doc.add_heading("Overall Class Distribution (%)", level=1)
    tbl = doc.add_table(rows=len(all_model_names) + 2, cols=len(CLASS_NAMES) + 2)
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    tbl.rows[0].cells[0].text = "Model"
    for j, c in enumerate(CLASS_NAMES, 1):
        tbl.rows[0].cells[j].text = c
    tbl.rows[0].cells[len(CLASS_NAMES) + 1].text = "Anomaly %"

    # Data rows
    for i, mn in enumerate(all_model_names, 1):
        counts = global_counts[mn]
        total = sum(counts.values())
        if total == 0:
            continue
        tbl.rows[i].cells[0].text = mn
        anomaly = (total - counts["Normal"]) / total * 100
        for j, c in enumerate(CLASS_NAMES, 1):
            pct = counts[c] / total * 100
            tbl.rows[i].cells[j].text = f"{pct:.1f}%"
        tbl.rows[i].cells[len(CLASS_NAMES) + 1].text = f"{anomaly:.1f}%"

    # Clinical row
    clin_row = len(all_model_names) + 1
    tbl.rows[clin_row].cells[0].text = "CLINICAL"
    for j, c in enumerate(CLASS_NAMES, 1):
        tbl.rows[clin_row].cells[j].text = f"{clinical[c]:.1f}%"
    tbl.rows[clin_row].cells[len(CLASS_NAMES) + 1].text = "98.0%"

    # MAE table in report
    doc.add_heading("MAE (Mean Absolute Error) vs Clinical", level=1)
    mae_tbl = doc.add_table(rows=len(mae_results) + 1, cols=len(CLASS_NAMES) + 2)
    mae_tbl.style = "Light Grid Accent 1"
    mae_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    mae_tbl.rows[0].cells[0].text = "Model"
    for j, c in enumerate(CLASS_NAMES, 1):
        mae_tbl.rows[0].cells[j].text = c
    mae_tbl.rows[0].cells[len(CLASS_NAMES) + 1].text = "MAE"

    for i, (mn, mae) in enumerate(ranked, 1):
        counts = global_counts[mn]
        total = sum(counts.values())
        mae_tbl.rows[i].cells[0].text = mn
        for j, c in enumerate(CLASS_NAMES, 1):
            predicted_pct = counts[c] / total * 100
            abs_err = abs(predicted_pct - clinical[c])
            mae_tbl.rows[i].cells[j].text = f"{abs_err:.1f}%"
        mae_tbl.rows[i].cells[len(CLASS_NAMES) + 1].text = f"{mae:.1f}%"

    # Per-frame summary (first 20 frames sample)
    doc.add_heading("Sample Per-Frame Breakdown (first 20)", level=1)
    sample_frames = frame_stats[:20]
    for fs in sample_frames:
        p = doc.add_paragraph()
        p.add_run(f"Frame {fs['frame_idx']} ({fs['n_sperm']} sperm): ").bold = True
        parts = []
        for mn in all_model_names[:3]:  # Show only first 3 models to save space
            counts = fs["counts"][mn]
            total_f = sum(counts.values())
            if total_f == 0:
                continue
            top_cls = max(counts, key=counts.get)
            top_pct = counts[top_cls] / total_f * 100
            parts.append(f"{mn}→{top_cls.split('_')[0]}({top_pct:.0f}%)")
        p.add_run(" | ".join(parts))

    doc.save(str(output_docx))
    print(f"\n📄 Report: {output_docx}")
    print(f"{'═'*70}")


if __name__ == "__main__":
    main()
