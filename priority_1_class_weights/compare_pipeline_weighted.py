"""
═══════════════════════════════════════════════════════════════════════════════
  Priority 1: Dual-Model Video Comparison Pipeline (Weighted Models)
  ──────────────────────────────────────────────────────────────────
  Same as the root compare_pipeline.py, but uses the WEIGHTED models
  trained with class-weighted loss from Priority 1.

  Usage:
      python priority_1_class_weights/compare_pipeline_weighted.py --video path/to/video.mp4
      python priority_1_class_weights/compare_pipeline_weighted.py --video path/to/video.mp4 --frames 5
═══════════════════════════════════════════════════════════════════════════════
"""

import cv2
import torch
import torch.nn as nn
import numpy as np
from torchvision import transforms, models
from ultralytics import YOLO
from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random
import argparse
import datetime
import sys

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))
from sperm_pipeline.pipeline import SpermAnalysisPipeline

# ═════════════════════════════════════════════════════════════════════════════
# CONFIGURATION
# ═════════════════════════════════════════════════════════════════════════════
BASE_DIR = Path(r"D:\Krishna\SMDMSSdataset")

# Detection model (shared — same for both)
DETECTION_MODEL = BASE_DIR / "sperm.v1i.yolov8" / "runs" / "detect" / "sperm_detection" / "weights" / "best.pt"

# ── Priority 1 Weighted Models ──
YOLO_CLS_WEIGHTED = BASE_DIR / "priority_1_class_weights" / "yolo_runs" / "sperm_cls_weighted" / "weights" / "best.pt"
EFFNET_WEIGHTED = BASE_DIR / "priority_1_class_weights" / "weights" / "efficientnet_b0_weighted.pt"

# ── Original (Unweighted) Models ──
YOLO_CLS_ORIGINAL = BASE_DIR / "runs" / "classify" / "runs" / "sperm_cls_balanced" / "weights" / "best.pt"
EFFNET_ORIGINAL = BASE_DIR / "efficientnet_training" / "weights" / "efficientnet_b0_best.pt"

IMG_SIZE = 224
NUM_FRAMES = 2

CLASS_NAMES = ["Combined_Anomaly", "Head_Anomaly", "Midpiece_Anomaly", "Normal", "Tail_Anomaly"]

OUTPUT_DOCX = BASE_DIR / "priority_1_class_weights" / "weighted_comparison_report.docx"


# ═════════════════════════════════════════════════════════════════════════════
# MODEL LOADING
# ═════════════════════════════════════════════════════════════════════════════

def load_efficientnet(weight_path, device, label=""):
    """Load an EfficientNet-B0 model from given weights."""
    print(f"  📦 Loading EfficientNet-B0 ({label})...")
    model = models.efficientnet_b0(weights=None)
    num_ftrs = model.classifier[1].in_features

    # Try with dropout classifier first (weighted model), fallback to simple
    try:
        model.classifier = nn.Sequential(
            nn.Dropout(p=0.4, inplace=True),
            nn.Linear(num_ftrs, len(CLASS_NAMES))
        )
        model.load_state_dict(torch.load(weight_path, map_location=device))
    except RuntimeError:
        model = models.efficientnet_b0(weights=None)
        num_ftrs = model.classifier[1].in_features
        model.classifier[1] = nn.Linear(num_ftrs, len(CLASS_NAMES))
        model.load_state_dict(torch.load(weight_path, map_location=device))

    model = model.to(device)
    model.eval()
    print(f"     ✓ Loaded: {weight_path.name}")
    return model


def load_yolo_cls(weight_path, label=""):
    """Load a YOLOv8 classifier."""
    print(f"  📦 Loading YOLOv8 ({label})...")
    model = YOLO(str(weight_path))
    print(f"     ✓ Loaded: {weight_path.name}")
    return model


# ═════════════════════════════════════════════════════════════════════════════
# INFERENCE
# ═════════════════════════════════════════════════════════════════════════════

def classify_effnet(model, crop_bgr, device):
    transform = transforms.Compose([
        transforms.Resize((IMG_SIZE, IMG_SIZE)),
        transforms.ToTensor(),
        transforms.Normalize([0.485, 0.456, 0.406], [0.229, 0.224, 0.225])
    ])
    crop_rgb = cv2.cvtColor(crop_bgr, cv2.COLOR_BGR2RGB)
    pil_img = Image.fromarray(crop_rgb)
    img_tensor = transform(pil_img).unsqueeze(0).to(device)

    with torch.no_grad():
        outputs = model(img_tensor)
        probs = torch.nn.functional.softmax(outputs, dim=1)
        confidence, pred_idx = torch.max(probs, 1)

    return CLASS_NAMES[pred_idx.item()], confidence.item()


def classify_yolo(model, crop_bgr):
    crop_resized = cv2.resize(crop_bgr, (IMG_SIZE, IMG_SIZE))
    results = model.predict(source=crop_resized, imgsz=IMG_SIZE, verbose=False)
    probs = results[0].probs
    pred_idx = probs.top1
    confidence = float(probs.top1conf)
    pred_name = results[0].names[pred_idx]
    return pred_name, confidence


# ═════════════════════════════════════════════════════════════════════════════
# VIDEO FRAME EXTRACTION
# ═════════════════════════════════════════════════════════════════════════════

def extract_random_frames(video_path, num_frames):
    cap = cv2.VideoCapture(str(video_path))
    if not cap.isOpened():
        print(f"❌ Cannot open video: {video_path}")
        return []

    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    fps = cap.get(cv2.CAP_PROP_FPS)
    print(f"  Total frames: {total_frames}, FPS: {fps}")

    if total_frames == 0:
        cap.release()
        return []

    n = min(num_frames, total_frames)
    indices = sorted(random.sample(range(total_frames), n))
    print(f"  Selected frames: {indices}")

    frames = []
    for idx in indices:
        cap.set(cv2.CAP_PROP_POS_FRAMES, idx)
        ret, frame = cap.read()
        if ret:
            frames.append((idx, frame))

    cap.release()
    return frames


# ═════════════════════════════════════════════════════════════════════════════
# REPORT GENERATION
# ═════════════════════════════════════════════════════════════════════════════

def create_report(results, video_path, frame_indices):
    doc = Document()

    title = doc.add_heading("Priority 1: Weighted vs Original Model Comparison", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Video: {Path(video_path).name}")
    doc.add_paragraph(f"Frames sampled: {frame_indices}")
    doc.add_paragraph(f"Total sperm crops: {len(results)}")

    # ── Summary: 4 models compared ──
    doc.add_heading("Model Performance Summary", level=1)

    total = len(results)
    if total == 0:
        doc.add_paragraph("No sperm detected.")
        doc.save(str(OUTPUT_DOCX))
        return

    models_info = [
        ("YOLOv8 Original", "yolo_orig_pred", "yolo_orig_conf"),
        ("YOLOv8 Weighted", "yolo_wt_pred", "yolo_wt_conf"),
        ("EffNet Original", "effnet_orig_pred", "effnet_orig_conf"),
        ("EffNet Weighted", "effnet_wt_pred", "effnet_wt_conf"),
    ]

    # Class distribution per model
    doc.add_heading("Class Distribution by Model", level=2)

    all_classes = sorted(CLASS_NAMES)
    dist = doc.add_table(rows=len(all_classes) + 1, cols=5)
    dist.style = "Light Grid Accent 1"
    dist.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Class"] + [m[0] for m in models_info]
    for j, h in enumerate(headers):
        dist.rows[0].cells[j].text = h
        for run in dist.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True

    for i, cls in enumerate(all_classes, start=1):
        dist.rows[i].cells[0].text = cls
        for j, (_, pred_key, _) in enumerate(models_info, start=1):
            count = sum(1 for r in results if r[pred_key] == cls)
            pct = count / total
            dist.rows[i].cells[j].text = f"{count} ({pct:.0%})"

    # Avg confidence per model
    doc.add_heading("Average Confidence by Model", level=2)

    conf_table = doc.add_table(rows=2, cols=5)
    conf_table.style = "Light Grid Accent 1"
    conf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    conf_table.rows[0].cells[0].text = "Metric"
    conf_table.rows[1].cells[0].text = "Avg Confidence"
    for run in conf_table.rows[0].cells[0].paragraphs[0].runs:
        run.bold = True

    for j, (name, _, conf_key) in enumerate(models_info, start=1):
        conf_table.rows[0].cells[j].text = name
        for run in conf_table.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True
        avg = sum(r[conf_key] for r in results) / total
        conf_table.rows[1].cells[j].text = f"{avg:.1%}"

    # ── Detailed per-crop comparison (Original vs Weighted) ──
    doc.add_heading("Per-Crop: Original vs Weighted", level=1)

    detail_headers = ["#", "Frame", "Sperm",
                      "YOLO Orig", "YOLO Wt", "YOLO Changed?",
                      "EffNet Orig", "EffNet Wt", "EffNet Changed?"]
    dtbl = doc.add_table(rows=len(results) + 1, cols=len(detail_headers))
    dtbl.style = "Light Grid Accent 1"
    dtbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(detail_headers):
        dtbl.rows[0].cells[j].text = h
        for run in dtbl.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True

    for i, r in enumerate(results, start=1):
        row = dtbl.rows[i]
        row.cells[0].text = str(i)
        row.cells[1].text = str(r["frame_idx"])
        row.cells[2].text = f"#{r['sperm_id']}"
        row.cells[3].text = f"{r['yolo_orig_pred']} ({r['yolo_orig_conf']:.0%})"
        row.cells[4].text = f"{r['yolo_wt_pred']} ({r['yolo_wt_conf']:.0%})"
        row.cells[5].text = "✅ Same" if r['yolo_orig_pred'] == r['yolo_wt_pred'] else "🔄 Changed"
        row.cells[6].text = f"{r['effnet_orig_pred']} ({r['effnet_orig_conf']:.0%})"
        row.cells[7].text = f"{r['effnet_wt_pred']} ({r['effnet_wt_conf']:.0%})"
        row.cells[8].text = "✅ Same" if r['effnet_orig_pred'] == r['effnet_wt_pred'] else "🔄 Changed"

    # ── Improvement Summary ──
    doc.add_heading("Improvement Summary", level=1)

    yolo_changed = sum(1 for r in results if r['yolo_orig_pred'] != r['yolo_wt_pred'])
    effnet_changed = sum(1 for r in results if r['effnet_orig_pred'] != r['effnet_wt_pred'])

    doc.add_paragraph(f"YOLOv8: {yolo_changed}/{total} predictions changed after weighting ({yolo_changed/total:.1%})")
    doc.add_paragraph(f"EfficientNet: {effnet_changed}/{total} predictions changed after weighting ({effnet_changed/total:.1%})")

    # Agreement rates
    orig_agree = sum(1 for r in results if r['yolo_orig_pred'] == r['effnet_orig_pred'])
    wt_agree = sum(1 for r in results if r['yolo_wt_pred'] == r['effnet_wt_pred'])
    doc.add_paragraph(f"\nOriginal models agreement: {orig_agree}/{total} ({orig_agree/total:.1%})")
    doc.add_paragraph(f"Weighted models agreement: {wt_agree}/{total} ({wt_agree/total:.1%})")

    doc.save(str(OUTPUT_DOCX))
    print(f"\n✅ Report saved: {OUTPUT_DOCX}")


# ═════════════════════════════════════════════════════════════════════════════
# MAIN
# ═════════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description="Priority 1: Weighted Model Video Comparison")
    parser.add_argument("--video", type=str, required=True, help="Path to input video")
    parser.add_argument("--frames", type=int, default=NUM_FRAMES, help="Random frames to extract")
    args = parser.parse_args()

    video_path = Path(args.video)
    if not video_path.exists():
        print(f"❌ Video not found: {video_path}")
        sys.exit(1)

    print("═" * 65)
    print("  Priority 1: Weighted vs Original Model Comparison")
    print("═" * 65)

    device = torch.device("cuda:0" if torch.cuda.is_available() else "cpu")
    print(f"  Device: {device}\n")

    # ── Extract frames ──
    print(f"🎥 Extracting {args.frames} random frames from: {video_path.name}")
    frames = extract_random_frames(video_path, args.frames)
    if not frames:
        sys.exit(1)

    # ── Load detection pipeline ──
    print("\n🔬 Loading sperm detection pipeline...")
    det_pipeline = SpermAnalysisPipeline(
        detection_model=str(DETECTION_MODEL),
        classification_model=str(YOLO_CLS_ORIGINAL)
    )
    if not det_pipeline.load_models():
        sys.exit(1)

    # ── Load all 4 classifiers ──
    print("\n🧠 Loading all classifiers...")

    # Check which models exist
    models_loaded = {}

    if YOLO_CLS_ORIGINAL.exists():
        models_loaded['yolo_orig'] = load_yolo_cls(YOLO_CLS_ORIGINAL, "Original")
    else:
        print(f"  ⚠️ YOLOv8 Original not found, skipping")

    if YOLO_CLS_WEIGHTED.exists():
        models_loaded['yolo_wt'] = load_yolo_cls(YOLO_CLS_WEIGHTED, "Weighted P1")
    else:
        print(f"  ⚠️ YOLOv8 Weighted not found, skipping")

    if EFFNET_ORIGINAL.exists():
        models_loaded['effnet_orig'] = load_efficientnet(EFFNET_ORIGINAL, device, "Original")
    else:
        print(f"  ⚠️ EfficientNet Original not found, skipping")

    if EFFNET_WEIGHTED.exists():
        models_loaded['effnet_wt'] = load_efficientnet(EFFNET_WEIGHTED, device, "Weighted P1")
    else:
        print(f"  ⚠️ EfficientNet Weighted not found, skipping")

    # ── Process frames ──
    all_results = []
    frame_indices = []

    for frame_idx, frame in frames:
        frame_indices.append(frame_idx)
        print(f"\n── Frame {frame_idx} ──")

        detections = det_pipeline.detect_sperm(frame)
        print(f"   Detected {len(detections)} sperm")

        for sid, det in enumerate(detections, start=1):
            crop = det_pipeline.crop_detection(frame, det["bbox"])
            if crop.size == 0:
                continue

            result = {"frame_idx": frame_idx, "sperm_id": sid}

            # Run all available classifiers
            if 'yolo_orig' in models_loaded:
                p, c = classify_yolo(models_loaded['yolo_orig'], crop)
                result["yolo_orig_pred"], result["yolo_orig_conf"] = p, c
            else:
                result["yolo_orig_pred"], result["yolo_orig_conf"] = "N/A", 0

            if 'yolo_wt' in models_loaded:
                p, c = classify_yolo(models_loaded['yolo_wt'], crop)
                result["yolo_wt_pred"], result["yolo_wt_conf"] = p, c
            else:
                result["yolo_wt_pred"], result["yolo_wt_conf"] = "N/A", 0

            if 'effnet_orig' in models_loaded:
                p, c = classify_effnet(models_loaded['effnet_orig'], crop, device)
                result["effnet_orig_pred"], result["effnet_orig_conf"] = p, c
            else:
                result["effnet_orig_pred"], result["effnet_orig_conf"] = "N/A", 0

            if 'effnet_wt' in models_loaded:
                p, c = classify_effnet(models_loaded['effnet_wt'], crop, device)
                result["effnet_wt_pred"], result["effnet_wt_conf"] = p, c
            else:
                result["effnet_wt_pred"], result["effnet_wt_conf"] = "N/A", 0

            # Print comparison
            yo = result["yolo_orig_pred"]
            yw = result["yolo_wt_pred"]
            eo = result["effnet_orig_pred"]
            ew = result["effnet_wt_pred"]
            y_icon = "✅" if yo == yw else "🔄"
            e_icon = "✅" if eo == ew else "🔄"

            print(f"   #{sid:2d}  YOLO: {yo:20s}→{yw:20s}{y_icon}  "
                  f"EffNet: {eo:20s}→{ew:20s}{e_icon}")

            all_results.append(result)

    # ── Generate report ──
    print(f"\n📝 Generating report... ({len(all_results)} crops)")
    create_report(all_results, video_path, frame_indices)

    # Console summary
    if all_results:
        total = len(all_results)
        yc = sum(1 for r in all_results if r['yolo_orig_pred'] != r['yolo_wt_pred'])
        ec = sum(1 for r in all_results if r['effnet_orig_pred'] != r['effnet_wt_pred'])
        print(f"\n{'═' * 65}")
        print(f"  YOLOv8:      {yc}/{total} predictions changed ({yc/total:.1%})")
        print(f"  EfficientNet: {ec}/{total} predictions changed ({ec/total:.1%})")
        print(f"  Report:       {OUTPUT_DOCX}")
        print(f"{'═' * 65}")


if __name__ == "__main__":
    main()
