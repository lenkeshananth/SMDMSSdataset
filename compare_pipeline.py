"""
═══════════════════════════════════════════════════════════════════════════════
  Dual-Model Video Comparison Pipeline
  ─────────────────────────────────────
  1. Accepts a video file as input
  2. Extracts 2 random frames from the video
  3. Detects sperm cells using YOLOv8 object detection
  4. Crops each detected sperm
  5. Classifies each crop using BOTH YOLOv8 (Balanced) and EfficientNet-B0
  6. Compares predictions side-by-side
  7. Generates a .docx report with full comparison

  Usage:
      python compare_pipeline.py --video path/to/video.mp4
      python compare_pipeline.py --video path/to/video.mp4 --frames 5
═══════════════════════════════════════════════════════════════════════════════
"""

import cv2
import torch
import torch.nn as nn
import numpy as np
from torchvision import transforms, models
from ultralytics import YOLO
from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random
import argparse
import datetime
import sys
import tempfile

# Add parent directory to path so we can import the pipeline
sys.path.insert(0, str(Path(__file__).parent))
from sperm_pipeline.pipeline import SpermAnalysisPipeline

# ═════════════════════════════════════════════════════════════════════════════
# CONFIGURATION
# ═════════════════════════════════════════════════════════════════════════════
BASE_DIR = Path(r"D:\Krishna\SMDMSSdataset")

# ── Model Paths ──
# Detection model (Stage 1 — shared by both)
DETECTION_MODEL = BASE_DIR / "sperm.v1i.yolov8" / "runs" / "detect" / "sperm_detection" / "weights" / "best.pt"

# Classification models (Stage 2 — compared head-to-head)
YOLO_CLS_MODEL = BASE_DIR / "runs" / "classify" / "runs" / "sperm_cls_balanced" / "weights" / "best.pt"
EFFICIENTNET_PATH = BASE_DIR / "efficientnet_training" / "weights" / "efficientnet_b0_best.pt"

# ── Settings ──
IMG_SIZE = 224
NUM_FRAMES = 2          # Number of random frames to extract
OUTPUT_DOCX = BASE_DIR / "dual_model_comparison_report.docx"

# ── Class names (must match training order) ──
CLASS_NAMES = ["Combined_Anomaly", "Head_Anomaly", "Midpiece_Anomaly", "Normal", "Tail_Anomaly"]


# ═════════════════════════════════════════════════════════════════════════════
# MODEL LOADING
# ═════════════════════════════════════════════════════════════════════════════

def load_efficientnet(device):
    """Load the trained EfficientNet-B0 classifier."""
    print("  📦 Loading EfficientNet-B0 classifier...")
    model = models.efficientnet_b0(weights=None)
    num_ftrs = model.classifier[1].in_features
    model.classifier[1] = nn.Linear(num_ftrs, len(CLASS_NAMES))
    model.load_state_dict(torch.load(EFFICIENTNET_PATH, map_location=device))
    model = model.to(device)
    model.eval()
    print(f"     ✓ Loaded from: {EFFICIENTNET_PATH}")
    return model


def load_yolo_classifier():
    """Load the trained YOLOv8 classification model."""
    print("  📦 Loading YOLOv8 (Balanced) classifier...")
    model = YOLO(str(YOLO_CLS_MODEL))
    print(f"     ✓ Loaded from: {YOLO_CLS_MODEL}")
    return model


# ═════════════════════════════════════════════════════════════════════════════
# INFERENCE HELPERS
# ═════════════════════════════════════════════════════════════════════════════

def classify_with_efficientnet(model, crop_bgr, device):
    """
    Classify a single BGR crop using EfficientNet-B0.
    Returns (predicted_class_name, confidence).
    """
    transform = transforms.Compose([
        transforms.Resize((IMG_SIZE, IMG_SIZE)),
        transforms.ToTensor(),
        transforms.Normalize([0.485, 0.456, 0.406], [0.229, 0.224, 0.225])
    ])
    # Convert BGR (OpenCV) → RGB (PIL)
    crop_rgb = cv2.cvtColor(crop_bgr, cv2.COLOR_BGR2RGB)
    pil_img = Image.fromarray(crop_rgb)
    img_tensor = transform(pil_img).unsqueeze(0).to(device)

    with torch.no_grad():
        outputs = model(img_tensor)
        probs = torch.nn.functional.softmax(outputs, dim=1)
        confidence, pred_idx = torch.max(probs, 1)

    return CLASS_NAMES[pred_idx.item()], confidence.item()


def classify_with_yolo(model, crop_bgr):
    """
    Classify a single BGR crop using YOLOv8 classification model.
    Returns (predicted_class_name, confidence).
    """
    crop_resized = cv2.resize(crop_bgr, (IMG_SIZE, IMG_SIZE))
    results = model.predict(source=crop_resized, imgsz=IMG_SIZE, verbose=False)
    probs = results[0].probs
    pred_idx = probs.top1
    confidence = float(probs.top1conf)
    pred_name = results[0].names[pred_idx]
    return pred_name, confidence


# ═════════════════════════════════════════════════════════════════════════════
# VIDEO FRAME EXTRACTION & DETECTION
# ═════════════════════════════════════════════════════════════════════════════

def extract_random_frames(video_path, num_frames):
    """Open a video, pick N random frame indices, and return the frames."""
    cap = cv2.VideoCapture(str(video_path))
    if not cap.isOpened():
        print(f"❌ Cannot open video: {video_path}")
        return []

    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    fps = cap.get(cv2.CAP_PROP_FPS)
    width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
    height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))

    print(f"  Total frames: {total_frames}")
    print(f"  FPS: {fps}")
    print(f"  Resolution: {width}x{height}")

    if total_frames == 0:
        cap.release()
        return []

    n = min(num_frames, total_frames)
    indices = sorted(random.sample(range(total_frames), n))
    print(f"  Selected random frame indices: {indices}")

    frames = []
    for idx in indices:
        cap.set(cv2.CAP_PROP_POS_FRAMES, idx)
        ret, frame = cap.read()
        if ret:
            frames.append((idx, frame))
        else:
            print(f"  ⚠️ Failed to read frame {idx}")

    cap.release()
    return frames


# ═════════════════════════════════════════════════════════════════════════════
# REPORT GENERATION
# ═════════════════════════════════════════════════════════════════════════════

def save_crop_image(crop_bgr, idx):
    """Save a crop to a temp file and return the path (for embedding in docx)."""
    tmp_path = Path(tempfile.gettempdir()) / f"crop_{idx}.png"
    cv2.imwrite(str(tmp_path), crop_bgr)
    return tmp_path


def create_report(results, video_path, frame_indices):
    """Generate the .docx comparison report."""
    doc = Document()

    # ── Title ──
    title = doc.add_heading("Dual-Model Video Comparison Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Video: {Path(video_path).name}")
    doc.add_paragraph(f"Frames sampled: {frame_indices}")
    doc.add_paragraph(f"Total sperm crops analyzed: {len(results)}")
    doc.add_paragraph(f"Models: YOLOv8 (Balanced) vs EfficientNet-B0")

    # ── Summary Statistics ──
    doc.add_heading("Summary Statistics", level=1)

    total = len(results)
    if total == 0:
        doc.add_paragraph("No sperm were detected in the sampled frames.")
        doc.save(str(OUTPUT_DOCX))
        return

    agree_count = sum(1 for r in results if r["yolo_pred"] == r["effnet_pred"])
    avg_yolo_conf = sum(r["yolo_conf"] for r in results) / total
    avg_effnet_conf = sum(r["effnet_conf"] for r in results) / total

    # Count per-class for each model
    yolo_class_counts = {}
    effnet_class_counts = {}
    for r in results:
        yolo_class_counts[r["yolo_pred"]] = yolo_class_counts.get(r["yolo_pred"], 0) + 1
        effnet_class_counts[r["effnet_pred"]] = effnet_class_counts.get(r["effnet_pred"], 0) + 1

    summary_data = [
        ("Metric", "Value"),
        ("Total Sperm Crops", str(total)),
        ("Models Agree", f"{agree_count}/{total} ({agree_count/total:.1%})"),
        ("Models Disagree", f"{total - agree_count}/{total} ({(total - agree_count)/total:.1%})"),
        ("Avg YOLOv8 Confidence", f"{avg_yolo_conf:.2%}"),
        ("Avg EfficientNet Confidence", f"{avg_effnet_conf:.2%}"),
    ]

    tbl = doc.add_table(rows=len(summary_data), cols=2)
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(summary_data):
        tbl.rows[i].cells[0].text = label
        tbl.rows[i].cells[1].text = value
        if i == 0:
            for cell in tbl.rows[i].cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True

    # ── Class Distribution Comparison ──
    doc.add_heading("Class Distribution Comparison", level=1)

    all_classes = sorted(set(list(yolo_class_counts.keys()) + list(effnet_class_counts.keys())))
    dist_table = doc.add_table(rows=len(all_classes) + 1, cols=3)
    dist_table.style = "Light Grid Accent 1"
    dist_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    dist_headers = ["Class", "YOLOv8 Count", "EfficientNet Count"]
    for j, h in enumerate(dist_headers):
        dist_table.rows[0].cells[j].text = h
        for run in dist_table.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True

    for i, cls in enumerate(all_classes, start=1):
        dist_table.rows[i].cells[0].text = cls
        dist_table.rows[i].cells[1].text = str(yolo_class_counts.get(cls, 0))
        dist_table.rows[i].cells[2].text = str(effnet_class_counts.get(cls, 0))

    # ── Detailed Per-Crop Comparison ──
    doc.add_heading("Detailed Per-Crop Comparison", level=1)

    detail_headers = ["#", "Frame", "Sperm", "YOLOv8 Pred", "YOLOv8 Conf",
                       "EffNet Pred", "EffNet Conf", "Agreement"]
    detail_table = doc.add_table(rows=len(results) + 1, cols=len(detail_headers))
    detail_table.style = "Light Grid Accent 1"
    detail_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(detail_headers):
        detail_table.rows[0].cells[j].text = h
        for run in detail_table.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True

    for i, r in enumerate(results, start=1):
        row = detail_table.rows[i]
        row.cells[0].text = str(i)
        row.cells[1].text = str(r["frame_idx"])
        row.cells[2].text = f"Sperm #{r['sperm_id']}"
        row.cells[3].text = r["yolo_pred"]
        row.cells[4].text = f"{r['yolo_conf']:.1%}"
        row.cells[5].text = r["effnet_pred"]
        row.cells[6].text = f"{r['effnet_conf']:.1%}"

        if r["yolo_pred"] == r["effnet_pred"]:
            row.cells[7].text = "✅ Agree"
        else:
            row.cells[7].text = "❌ Disagree"

    # ── Disagreement Details ──
    disagreements = [r for r in results if r["yolo_pred"] != r["effnet_pred"]]
    if disagreements:
        doc.add_heading("Disagreement Details", level=1)
        doc.add_paragraph(
            f"The models disagreed on {len(disagreements)} out of {total} crops "
            f"({len(disagreements)/total:.1%})."
        )
        for idx, r in enumerate(disagreements, start=1):
            p = doc.add_paragraph()
            p.add_run(f"{idx}. Frame {r['frame_idx']}, Sperm #{r['sperm_id']}").bold = True
            p.add_run(f"\n   YOLOv8: {r['yolo_pred']} ({r['yolo_conf']:.1%})")
            p.add_run(f"\n   EfficientNet: {r['effnet_pred']} ({r['effnet_conf']:.1%})")

    # ── Save ──
    doc.save(str(OUTPUT_DOCX))
    print(f"\n✅ Report saved to: {OUTPUT_DOCX}")


# ═════════════════════════════════════════════════════════════════════════════
# MAIN PIPELINE
# ═════════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description="Dual-Model Video Comparison Pipeline")
    parser.add_argument("--video", type=str, required=True, help="Path to input video file")
    parser.add_argument("--frames", type=int, default=NUM_FRAMES,
                        help=f"Number of random frames to extract (default: {NUM_FRAMES})")
    args = parser.parse_args()

    video_path = Path(args.video)
    if not video_path.exists():
        print(f"❌ Video file not found: {video_path}")
        sys.exit(1)

    print("═" * 65)
    print("    Dual-Model Video Comparison Pipeline")
    print("═" * 65)

    device = torch.device("cuda:0" if torch.cuda.is_available() else "cpu")
    print(f"  Device: {device}\n")

    # ── Step 1: Extract random frames ──
    print(f"🎥 Extracting {args.frames} random frames from: {video_path.name}")
    frames = extract_random_frames(video_path, args.frames)

    if not frames:
        print("❌ No frames could be read from the video.")
        sys.exit(1)

    # ── Step 2: Load detection pipeline (for sperm cropping) ──
    print("\n🔬 Loading sperm detection pipeline...")
    detection_pipeline = SpermAnalysisPipeline(
        detection_model=str(DETECTION_MODEL),
        classification_model=str(YOLO_CLS_MODEL)  # needed for init but we'll classify manually
    )
    if not detection_pipeline.load_models():
        print("❌ Failed to load detection models.")
        sys.exit(1)

    # ── Step 3: Load both classifiers ──
    print("\n🧠 Loading classification models...")
    effnet_model = load_efficientnet(device)
    yolo_cls_model = load_yolo_classifier()

    # ── Step 4: Process each frame ──
    all_results = []
    frame_indices = []

    for frame_idx, frame in frames:
        frame_indices.append(frame_idx)
        print(f"\n── Frame {frame_idx} ──")

        # Detect sperm in the frame
        detections = detection_pipeline.detect_sperm(frame)
        print(f"   Detected {len(detections)} sperm cells")

        for sperm_id, det in enumerate(detections, start=1):
            bbox = det["bbox"]
            crop = detection_pipeline.crop_detection(frame, bbox)

            if crop.size == 0:
                continue

            # Classify with both models
            yolo_pred, yolo_conf = classify_with_yolo(yolo_cls_model, crop)
            effnet_pred, effnet_conf = classify_with_efficientnet(effnet_model, crop, device)

            match_icon = "✅" if yolo_pred == effnet_pred else "❌"
            print(f"   Sperm #{sperm_id:2d}  "
                  f"YOLO={yolo_pred:20s}({yolo_conf:.1%})  "
                  f"EffNet={effnet_pred:20s}({effnet_conf:.1%})  "
                  f"{match_icon}")

            all_results.append({
                "frame_idx": frame_idx,
                "sperm_id": sperm_id,
                "yolo_pred": yolo_pred,
                "yolo_conf": yolo_conf,
                "effnet_pred": effnet_pred,
                "effnet_conf": effnet_conf,
            })

    # ── Step 5: Generate report ──
    print(f"\n📝 Generating comparison report...")
    print(f"   Total crops compared: {len(all_results)}")
    create_report(all_results, video_path, frame_indices)

    # ── Quick console summary ──
    if all_results:
        agree = sum(1 for r in all_results if r["yolo_pred"] == r["effnet_pred"])
        total = len(all_results)
        print(f"\n{'═' * 65}")
        print(f"   Agreement: {agree}/{total} ({agree/total:.1%})")
        print(f"   Report:    {OUTPUT_DOCX}")
        print(f"{'═' * 65}")


if __name__ == "__main__":
    main()
