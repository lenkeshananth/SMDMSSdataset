"""
═══════════════════════════════════════════════════════════════════════════════
  Priority 4: Ensemble Model — Sperm Morphology Classification
  ─────────────────────────────────────────────────────────────
  Combines YOLOv8 and EfficientNet-B0 predictions by averaging their
  probability distributions. This exploits the different biases of each
  model to produce more robust classifications.

  - YOLOv8 tends to over-predict Normal
  - EfficientNet tends to over-predict Combined_Anomaly
  - Ensemble balances these biases out

  Usage:
      python priority_4_ensemble/ensemble_classifier.py --video path/to/video.mp4
      python priority_4_ensemble/ensemble_classifier.py --video path/to/video.mp4 --frames 5
═══════════════════════════════════════════════════════════════════════════════
"""

import cv2
import torch
import torch.nn as nn
import numpy as np
from torchvision import transforms, models
from ultralytics import YOLO
from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random
import argparse
import datetime
import sys

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))
from sperm_pipeline.pipeline import SpermAnalysisPipeline

# ═════════════════════════════════════════════════════════════════════════════
# CONFIGURATION
# ═════════════════════════════════════════════════════════════════════════════
BASE_DIR = Path(r"D:\Krishna\SMDMSSdataset")

# Detection model
DETECTION_MODEL = BASE_DIR / "sperm.v1i.yolov8" / "runs" / "detect" / "sperm_detection" / "weights" / "best.pt"

# Use the WEIGHTED (Priority 1) models for ensembling — they're the best individual models
YOLO_CLS_PATH = BASE_DIR / "priority_1_class_weights" / "yolo_runs" / "sperm_cls_weighted" / "weights" / "best.pt"
EFFNET_PATH = BASE_DIR / "priority_1_class_weights" / "weights" / "efficientnet_b0_weighted.pt"

# Fallback to original models if weighted don't exist
YOLO_CLS_FALLBACK = BASE_DIR / "runs" / "classify" / "runs" / "sperm_cls_balanced" / "weights" / "best.pt"
EFFNET_FALLBACK = BASE_DIR / "efficientnet_training" / "weights" / "efficientnet_b0_best.pt"

IMG_SIZE = 224
NUM_FRAMES = 2

CLASS_NAMES = ["Combined_Anomaly", "Head_Anomaly", "Midpiece_Anomaly", "Normal", "Tail_Anomaly"]

# ── Ensemble Weights ──
# How much to trust each model's predictions.
# EfficientNet gets slightly more weight because it's better at detecting anomalies.
YOLO_WEIGHT = 0.4
EFFNET_WEIGHT = 0.6

OUTPUT_DOCX = BASE_DIR / "priority_4_ensemble" / "ensemble_comparison_report.docx"


# ═════════════════════════════════════════════════════════════════════════════
# MODEL LOADING
# ═════════════════════════════════════════════════════════════════════════════

def load_efficientnet(device):
    """Load EfficientNet-B0 classifier."""
    # Try weighted first, fallback to original
    weight_path = EFFNET_PATH if EFFNET_PATH.exists() else EFFNET_FALLBACK
    print(f"  📦 Loading EfficientNet-B0...")
    print(f"     Using: {weight_path.parent.parent.name}/{weight_path.name}")

    model = models.efficientnet_b0(weights=None)
    num_ftrs = model.classifier[1].in_features

    # Try dropout-based classifier (weighted model), fallback to simple
    try:
        model.classifier = nn.Sequential(
            nn.Dropout(p=0.4, inplace=True),
            nn.Linear(num_ftrs, len(CLASS_NAMES))
        )
        model.load_state_dict(torch.load(weight_path, map_location=device))
    except RuntimeError:
        model = models.efficientnet_b0(weights=None)
        num_ftrs = model.classifier[1].in_features
        model.classifier[1] = nn.Linear(num_ftrs, len(CLASS_NAMES))
        model.load_state_dict(torch.load(weight_path, map_location=device))

    model = model.to(device)
    model.eval()
    print(f"     ✓ Loaded")
    return model


def load_yolo():
    """Load YOLOv8 classifier."""
    weight_path = YOLO_CLS_PATH if YOLO_CLS_PATH.exists() else YOLO_CLS_FALLBACK
    print(f"  📦 Loading YOLOv8 classifier...")
    print(f"     Using: {weight_path.parent.parent.name}/{weight_path.name}")
    model = YOLO(str(weight_path))
    print(f"     ✓ Loaded")
    return model


# ═════════════════════════════════════════════════════════════════════════════
# INFERENCE — FULL PROBABILITY VECTORS
# ═════════════════════════════════════════════════════════════════════════════

def get_effnet_probs(model, crop_bgr, device):
    """Get full probability distribution from EfficientNet."""
    transform = transforms.Compose([
        transforms.Resize((IMG_SIZE, IMG_SIZE)),
        transforms.ToTensor(),
        transforms.Normalize([0.485, 0.456, 0.406], [0.229, 0.224, 0.225])
    ])
    crop_rgb = cv2.cvtColor(crop_bgr, cv2.COLOR_BGR2RGB)
    pil_img = Image.fromarray(crop_rgb)
    img_tensor = transform(pil_img).unsqueeze(0).to(device)

    with torch.no_grad():
        outputs = model(img_tensor)
        probs = torch.nn.functional.softmax(outputs, dim=1)[0].cpu().numpy()

    return probs  # shape: (5,)


def get_yolo_probs(model, crop_bgr):
    """Get full probability distribution from YOLOv8."""
    crop_resized = cv2.resize(crop_bgr, (IMG_SIZE, IMG_SIZE))
    results = model.predict(source=crop_resized, imgsz=IMG_SIZE, verbose=False)
    probs_tensor = results[0].probs.data.cpu().numpy()

    # YOLOv8 class order may differ from our CLASS_NAMES — remap
    yolo_names = results[0].names
    remapped = np.zeros(len(CLASS_NAMES))
    for yolo_idx, name in yolo_names.items():
        if name in CLASS_NAMES:
            our_idx = CLASS_NAMES.index(name)
            remapped[our_idx] = probs_tensor[yolo_idx]

    return remapped  # shape: (5,)


def ensemble_predict(yolo_probs, effnet_probs):
    """
    Combine two probability vectors using weighted average.
    Returns (predicted_class_name, confidence, combined_probs).
    """
    combined = YOLO_WEIGHT * yolo_probs + EFFNET_WEIGHT * effnet_probs

    # Normalize to sum to 1
    combined = combined / combined.sum()

    pred_idx = np.argmax(combined)
    confidence = combined[pred_idx]

    return CLASS_NAMES[pred_idx], float(confidence), combined


# ═════════════════════════════════════════════════════════════════════════════
# VIDEO PROCESSING
# ═════════════════════════════════════════════════════════════════════════════

def extract_random_frames(video_path, num_frames):
    cap = cv2.VideoCapture(str(video_path))
    if not cap.isOpened():
        print(f"❌ Cannot open video: {video_path}")
        return []

    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    fps = cap.get(cv2.CAP_PROP_FPS)
    print(f"  Total frames: {total_frames}, FPS: {fps}")

    if total_frames == 0:
        cap.release()
        return []

    n = min(num_frames, total_frames)
    indices = sorted(random.sample(range(total_frames), n))
    print(f"  Selected frames: {indices}")

    frames = []
    for idx in indices:
        cap.set(cv2.CAP_PROP_POS_FRAMES, idx)
        ret, frame = cap.read()
        if ret:
            frames.append((idx, frame))

    cap.release()
    return frames


# ═════════════════════════════════════════════════════════════════════════════
# REPORT GENERATION
# ═════════════════════════════════════════════════════════════════════════════

def create_report(results, video_path, frame_indices):
    doc = Document()

    title = doc.add_heading("Priority 4: Ensemble Model Comparison Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Video: {Path(video_path).name}")
    doc.add_paragraph(f"Frames sampled: {frame_indices}")
    doc.add_paragraph(f"Total sperm crops: {len(results)}")
    doc.add_paragraph(f"Ensemble weights: YOLOv8 = {YOLO_WEIGHT}, EfficientNet = {EFFNET_WEIGHT}")

    total = len(results)
    if total == 0:
        doc.add_paragraph("No sperm detected.")
        doc.save(str(OUTPUT_DOCX))
        return

    # ── Class Distribution: Individual vs Ensemble ──
    doc.add_heading("Class Distribution: Individual vs Ensemble", level=1)

    model_keys = [
        ("YOLOv8", "yolo_pred"),
        ("EfficientNet", "effnet_pred"),
        ("Ensemble", "ensemble_pred"),
    ]

    dist = doc.add_table(rows=len(CLASS_NAMES) + 1, cols=4)
    dist.style = "Light Grid Accent 1"
    dist.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Class"] + [m[0] for m in model_keys]
    for j, h in enumerate(headers):
        dist.rows[0].cells[j].text = h
        for run in dist.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True

    for i, cls in enumerate(CLASS_NAMES, start=1):
        dist.rows[i].cells[0].text = cls
        for j, (_, key) in enumerate(model_keys, start=1):
            count = sum(1 for r in results if r[key] == cls)
            dist.rows[i].cells[j].text = f"{count} ({count/total:.0%})"

    # ── Confidence Comparison ──
    doc.add_heading("Average Confidence", level=1)

    conf_tbl = doc.add_table(rows=2, cols=4)
    conf_tbl.style = "Light Grid Accent 1"
    conf_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    conf_tbl.rows[0].cells[0].text = ""
    conf_tbl.rows[1].cells[0].text = "Avg Confidence"

    conf_keys = [
        ("YOLOv8", "yolo_conf"),
        ("EfficientNet", "effnet_conf"),
        ("Ensemble", "ensemble_conf"),
    ]
    for j, (name, key) in enumerate(conf_keys, start=1):
        conf_tbl.rows[0].cells[j].text = name
        for run in conf_tbl.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True
        avg = sum(r[key] for r in results) / total
        conf_tbl.rows[1].cells[j].text = f"{avg:.1%}"

    # ── Detailed Per-Crop Table ──
    doc.add_heading("Detailed Per-Crop Comparison", level=1)

    detail_headers = ["#", "Frame", "Sperm",
                      "YOLOv8", "EffNet", "ENSEMBLE",
                      "Ensemble Conf"]
    dtbl = doc.add_table(rows=total + 1, cols=len(detail_headers))
    dtbl.style = "Light Grid Accent 1"
    dtbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(detail_headers):
        dtbl.rows[0].cells[j].text = h
        for run in dtbl.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True

    for i, r in enumerate(results, start=1):
        row = dtbl.rows[i]
        row.cells[0].text = str(i)
        row.cells[1].text = str(r["frame_idx"])
        row.cells[2].text = f"#{r['sperm_id']}"
        row.cells[3].text = f"{r['yolo_pred']} ({r['yolo_conf']:.0%})"
        row.cells[4].text = f"{r['effnet_pred']} ({r['effnet_conf']:.0%})"
        row.cells[5].text = r["ensemble_pred"]
        row.cells[6].text = f"{r['ensemble_conf']:.1%}"

    # ── Where Ensemble Differs from Both ──
    doc.add_heading("Ensemble Tiebreaker Cases", level=1)
    doc.add_paragraph(
        "Cases where YOLOv8 and EfficientNet disagreed and the ensemble resolved the conflict:"
    )

    disagree_cases = [r for r in results if r['yolo_pred'] != r['effnet_pred']]
    if disagree_cases:
        for idx, r in enumerate(disagree_cases, start=1):
            p = doc.add_paragraph()
            p.add_run(f"{idx}. Frame {r['frame_idx']}, Sperm #{r['sperm_id']}").bold = True
            p.add_run(f"\n   YOLOv8: {r['yolo_pred']} ({r['yolo_conf']:.1%})")
            p.add_run(f"\n   EfficientNet: {r['effnet_pred']} ({r['effnet_conf']:.1%})")
            p.add_run(f"\n   → Ensemble chose: ").bold = True
            p.add_run(f"{r['ensemble_pred']} ({r['ensemble_conf']:.1%})")
    else:
        doc.add_paragraph("All models agreed — no tiebreaker cases.")

    # ── Agreement Stats ──
    doc.add_heading("Agreement Statistics", level=1)

    yolo_effnet_agree = sum(1 for r in results if r['yolo_pred'] == r['effnet_pred'])
    ens_yolo_agree = sum(1 for r in results if r['ensemble_pred'] == r['yolo_pred'])
    ens_effnet_agree = sum(1 for r in results if r['ensemble_pred'] == r['effnet_pred'])

    stats = [
        ("YOLOv8 ↔ EfficientNet agreement", yolo_effnet_agree),
        ("Ensemble sided with YOLOv8", ens_yolo_agree),
        ("Ensemble sided with EfficientNet", ens_effnet_agree),
    ]
    for label, count in stats:
        doc.add_paragraph(f"{label}: {count}/{total} ({count/total:.1%})")

    doc.save(str(OUTPUT_DOCX))
    print(f"\n✅ Report saved: {OUTPUT_DOCX}")


# ═════════════════════════════════════════════════════════════════════════════
# MAIN
# ═════════════════════════════════════════════════════════════════════════════

def main():
    global YOLO_WEIGHT, EFFNET_WEIGHT

    parser = argparse.ArgumentParser(description="Priority 4: Ensemble Model Video Comparison")
    parser.add_argument("--video", type=str, required=True, help="Path to input video")
    parser.add_argument("--frames", type=int, default=NUM_FRAMES, help="Random frames to extract")
    parser.add_argument("--yolo-weight", type=float, default=YOLO_WEIGHT,
                        help=f"YOLOv8 ensemble weight (default: {YOLO_WEIGHT})")
    parser.add_argument("--effnet-weight", type=float, default=EFFNET_WEIGHT,
                        help=f"EfficientNet ensemble weight (default: {EFFNET_WEIGHT})")
    args = parser.parse_args()

    YOLO_WEIGHT = args.yolo_weight
    EFFNET_WEIGHT = args.effnet_weight

    video_path = Path(args.video)
    if not video_path.exists():
        print(f"❌ Video not found: {video_path}")
        sys.exit(1)

    print("═" * 65)
    print("  Priority 4: Ensemble Model Classification")
    print("═" * 65)

    device = torch.device("cuda:0" if torch.cuda.is_available() else "cpu")
    print(f"  Device: {device}")
    print(f"  Ensemble: {YOLO_WEIGHT:.0%} YOLOv8 + {EFFNET_WEIGHT:.0%} EfficientNet\n")

    # ── Extract frames ──
    print(f"🎥 Extracting {args.frames} random frames from: {video_path.name}")
    frames = extract_random_frames(video_path, args.frames)
    if not frames:
        sys.exit(1)

    # ── Load detection pipeline ──
    print("\n🔬 Loading detection pipeline...")
    yolo_cls_path = YOLO_CLS_PATH if YOLO_CLS_PATH.exists() else YOLO_CLS_FALLBACK
    det_pipeline = SpermAnalysisPipeline(
        detection_model=str(DETECTION_MODEL),
        classification_model=str(yolo_cls_path)
    )
    if not det_pipeline.load_models():
        sys.exit(1)

    # ── Load classifiers ──
    print("\n🧠 Loading classifiers for ensemble...")
    effnet_model = load_efficientnet(device)
    yolo_model = load_yolo()

    # ── Process frames ──
    all_results = []
    frame_indices = []

    for frame_idx, frame in frames:
        frame_indices.append(frame_idx)
        print(f"\n── Frame {frame_idx} ──")

        detections = det_pipeline.detect_sperm(frame)
        print(f"   Detected {len(detections)} sperm")

        for sid, det in enumerate(detections, start=1):
            crop = det_pipeline.crop_detection(frame, det["bbox"])
            if crop.size == 0:
                continue

            # Get full probability distributions
            yolo_probs = get_yolo_probs(yolo_model, crop)
            effnet_probs = get_effnet_probs(effnet_model, crop, device)

            # Individual predictions
            yolo_idx = np.argmax(yolo_probs)
            yolo_pred = CLASS_NAMES[yolo_idx]
            yolo_conf = float(yolo_probs[yolo_idx])

            effnet_idx = np.argmax(effnet_probs)
            effnet_pred = CLASS_NAMES[effnet_idx]
            effnet_conf = float(effnet_probs[effnet_idx])

            # Ensemble prediction
            ens_pred, ens_conf, ens_probs = ensemble_predict(yolo_probs, effnet_probs)

            # Icons
            agree = "✅" if yolo_pred == effnet_pred else "❌"
            ens_side = ""
            if yolo_pred != effnet_pred:
                if ens_pred == yolo_pred:
                    ens_side = " →YOLO"
                elif ens_pred == effnet_pred:
                    ens_side = " →EffNet"
                else:
                    ens_side = " →NEW"

            print(f"   #{sid:2d}  Y={yolo_pred:20s}({yolo_conf:.0%})  "
                  f"E={effnet_pred:20s}({effnet_conf:.0%})  "
                  f"{agree}  ENS={ens_pred:20s}({ens_conf:.0%}){ens_side}")

            all_results.append({
                "frame_idx": frame_idx,
                "sperm_id": sid,
                "yolo_pred": yolo_pred,
                "yolo_conf": yolo_conf,
                "effnet_pred": effnet_pred,
                "effnet_conf": effnet_conf,
                "ensemble_pred": ens_pred,
                "ensemble_conf": ens_conf,
            })

    # ── Generate report ──
    print(f"\n📝 Generating report... ({len(all_results)} crops)")
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    create_report(all_results, video_path, frame_indices)

    # ── Console summary ──
    if all_results:
        total = len(all_results)
        agree = sum(1 for r in all_results if r['yolo_pred'] == r['effnet_pred'])
        ens_yolo = sum(1 for r in all_results if r['ensemble_pred'] == r['yolo_pred'])
        ens_effnet = sum(1 for r in all_results if r['ensemble_pred'] == r['effnet_pred'])

        print(f"\n{'═' * 65}")
        print(f"  Models agreed:                {agree}/{total} ({agree/total:.1%})")
        print(f"  Ensemble sided with YOLOv8:   {ens_yolo}/{total} ({ens_yolo/total:.1%})")
        print(f"  Ensemble sided with EffNet:   {ens_effnet}/{total} ({ens_effnet/total:.1%})")

        # Class distribution
        print(f"\n  Ensemble class distribution:")
        for cls in CLASS_NAMES:
            count = sum(1 for r in all_results if r['ensemble_pred'] == cls)
            print(f"    {cls:25s} {count:3d} ({count/total:.0%})")

        print(f"\n  Report: {OUTPUT_DOCX}")
        print(f"{'═' * 65}")


if __name__ == "__main__":
    main()
